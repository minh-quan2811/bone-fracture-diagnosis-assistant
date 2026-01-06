from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches
import io
from datetime import datetime


def draw_bounding_boxes(image, student_bbox, model_bbox):
    """
    Draw student (blue) and model (red) bounding boxes on image
    """
    img_copy = image.copy()
    draw = ImageDraw.Draw(img_copy)
    
    # Draw student bbox (blue)
    draw.rectangle(
        [student_bbox["x_min"], student_bbox["y_min"], 
         student_bbox["x_max"], student_bbox["y_max"]],
        outline="blue",
        width=3
    )
    
    # Draw model bbox (red)
    draw.rectangle(
        [model_bbox["x_min"], model_bbox["y_min"], 
         model_bbox["x_max"], model_bbox["y_max"]],
        outline="red",
        width=3
    )
    
    return img_copy


def create_evaluation_report(
    image_with_boxes,
    student_class,
    model_class,
    student_note,
    evaluation_text,
    output_path
):
    """
    Create a .docx report with evaluation results
    
    Args:
        image_with_boxes: PIL Image with bounding boxes
        student_class: Student's predicted fracture type
        model_class: Model's predicted fracture type
        student_note: Student's reasoning
        evaluation_text: LLM evaluation feedback
        output_path: Path to save .docx file
    """
    doc = Document()
    
    # Title
    doc.add_heading('Fracture Classification Evaluation Report', 0)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Student Prediction Section
    doc.add_heading('Student Prediction', level=1)
    doc.add_paragraph(f'Predicted Fracture Type: {student_class}')
    doc.add_paragraph(f'Student Note: {student_note}')
    
    # Model Prediction Section
    doc.add_heading('Model Prediction', level=1)
    doc.add_paragraph(f'Predicted Fracture Type: {model_class}')
    
    # Image with Bounding Boxes
    doc.add_heading('Annotated X-ray Image', level=1)
    doc.add_paragraph('Blue Box: Student Annotation | Red Box: Model Prediction')
    
    # Save image to bytes and add to document
    img_byte_arr = io.BytesIO()
    image_with_boxes.save(img_byte_arr, format='PNG')
    img_byte_arr.seek(0)
    doc.add_picture(img_byte_arr, width=Inches(5))
    
    # Evaluation Section
    doc.add_heading('Instructor Evaluation', level=1)
    doc.add_paragraph(evaluation_text)
    
    # Save document
    doc.save(output_path)
    print(f'Report saved to: {output_path}')